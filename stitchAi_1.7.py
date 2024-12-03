from flask import Flask, request, jsonify, render_template_string, send_from_directory
from werkzeug.utils import secure_filename
import os
import requests
import concurrent.futures
from openai import OpenAI
from docx import Document  # 导入python-docx库

app = Flask(__name__)

# 设置文件上传的目录
UPLOAD_FOLDER = 'uploads/'
DOWNLOAD_FOLDER = 'downloads/'  # 新增下载文件夹
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

class ErnieBotService:
    def __init__(self, api_key, api_endpoint):
        self.api_key = api_key
        self.api_endpoint = api_endpoint

    def chatMessage(self, message):
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.api_key}'
        }
        data = {
            'query': message,
        }
        response = requests.post(self.api_endpoint, headers=headers, json=data)
        if response.status_code == 200:
            return response.json().get('answer', '未能从文心一言API获取答案。')
        else:
            return f'文心一言API调用失败，状态码：{response.status_code}'

class KimiService:
    def __init__(self, api_key, base_url, max_tokens=1024):  # 增加了max_tokens参数
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.max_tokens = max_tokens
    
    def chatMessage(self, message):
        try:
            completion = self.client.chat.completions.create(
                model="moonshot-v1-8k",
                messages=[{"role": "user", "content": message}],
                temperature=0.3,
                max_tokens=self.max_tokens  # 设置更大的max_tokens值
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"调用Kimi API时出错：{e}")
            return None

class AIAssistantDispatcher:
    def __init__(self):
        self.assistants = {
            'ernie': ErnieBotService(api_key="XVdKBjs4gqW15WcrnEAHRIQk", api_endpoint="https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/completions_pro?access_token=24.77d1a835729d0197520837d72f489451.2592000.1733388232.282335-116123794"),
            'kimi': KimiService(api_key="sk-P836YBvTryUCPHMSreuFMYxMA6wPQEckSe592YC3aOfWcWHU", base_url="https://api.moonshot.cn/v1")
        }
    
    def get_best_response(self, message, file_path=None):
        if file_path:
            # 如果有文件，就将文件内容作为消息的一部分
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    message += " " + content
            except Exception as e:
                return f"文件读取失败：{e}"

        # 使用并行请求多个AI助手生成不同部分的回答
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = {executor.submit(assistant.chatMessage, message): name for name, assistant in self.assistants.items()}
            responses = {}
            for future in concurrent.futures.as_completed(futures):
                name = futures[future]
                try:
                    response = future.result()
                    if response:
                        responses[name] = response
                except Exception as e:
                    responses[name] = f"调用{futures[future]}时发生错误：{e}"

        # 合并不同模型的回答
        merged_response = "\n".join(responses.values())

        # 如果生成的回答小于100个字符，则增加扩展内容
        if len(merged_response) < 100:
            merged_response += " 这是为了增加回答的长度而添加的额外内容。"

        return merged_response if merged_response else "对不起，我无法回答您的问题。"
    
dispatcher = AIAssistantDispatcher()

# 定义默认路由，显示上传和提问的表单
@app.route('/')
def home():
    return render_template_string('''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>AI Assistant</title>
        <!-- 引入Bootstrap CSS -->
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
    </head>
    <body>
        <div class="container mt-5">
            <h1 class="text-center">AI Assistant</h1>
            <div class="row justify-content-center">
                <div class="col-md-6">
                    <form action="/upload" method="post" enctype="multipart/form-data" class="mb-3">
                        <div class="input-group">
                            <div class="custom-file">
                                <input type="file" class="custom-file-input" id="inputGroupFile01" name="file" required>
                                <label class="custom-file-label" for="inputGroupFile01">Choose file</label>
                            </div>
                            <div class="input-group-append">
                                <input type="submit" class="btn btn-primary" value="Upload and解读 File">
                            </div>
                        </div>
                    </form>
                    <form action="/ask" method="post" class="mb-3">
                        <div class="input-group">
                            <input type="text" class="form-control" placeholder="Enter your question" name="question" required>
                            <div class="input-group-append">
                                <input type="submit" class="btn btn-outline-secondary" value="Ask a Question">
                            </div>
                        </div>
                    </form>
                </div>
            </div>
        </div>
        <!-- 引入Bootstrap JS -->
        <script src="https://code.jquery.com/jquery-3.5.1.slim.min.js"></script>
        <script src="https://cdn.jsdelivr.net/npm/@popperjs/core@2.9.3/dist/umd/popper.min.js"></script>
        <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/js/bootstrap.min.js"></script>
    </body>
    </html>
    ''')

# 文件上传路由
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        # 上传后直接解读文件
        best_answer = dispatcher.get_best_response("", file_path)
        
        # 将答案写入txt文件
        txt_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{filename}.txt')
        with open(txt_file_path, 'w', encoding='utf-8') as f:
            f.write(best_answer)
        
        # 也可以将答案写入docx文件（可选）
        docx_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{filename}.docx')
        doc = Document()
        doc.add_paragraph(best_answer)
        doc.save(docx_file_path)
        
        # 返回下载链接（或者你可以直接返回答案和文件路径）
        return jsonify({
            'answer': best_answer,
            'txt_download_url': f'/downloads/{os.path.basename(txt_file_path)}',
            'docx_download_url': f'/downloads/{os.path.basename(docx_file_path)}'
        })

# 问题回答路由
@app.route('/ask', methods=['POST'])
def ask_question():
    user_question = request.form.get('question')
    if not user_question:
        return jsonify({'error': 'No question provided'}), 400
    best_answer = dispatcher.get_best_response(user_question)
    
    # 将答案写入txt文件
    txt_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{user_question.replace(" ", "_")}.txt')
    with open(txt_file_path, 'w', encoding='utf-8') as f:
        f.write(best_answer)
    
    # 也可以将答案写入docx文件（可选）
    docx_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{user_question.replace(" ", "_")}.docx')
    doc = Document()
    doc.add_paragraph(best_answer)
    doc.save(docx_file_path)
    
    # 返回下载链接（或者你可以直接返回答案和文件路径）
    return jsonify({
        'answer': best_answer,
        'txt_download_url': f'/downloads/{os.path.basename(txt_file_path)}',
        'docx_download_url': f'/downloads/{os.path.basename(docx_file_path)}'
    })

# 新增下载路由
@app.route('/downloads/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(app.config['DOWNLOAD_FOLDER'], filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
