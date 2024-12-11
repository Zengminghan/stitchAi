from flask import Flask, request, jsonify, render_template_string, send_from_directory
from werkzeug.utils import secure_filename
import os
import requests
import concurrent.futures
from openai import OpenAI
from docx import Document  # 导入python-docx库
import torch
from transformers import BertTokenizer, BertModel
import json

app = Flask(__name__)

# 设置文件上传的目录
UPLOAD_FOLDER = 'uploads/'
DOWNLOAD_FOLDER = 'downloads/'  # 新增下载文件夹
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

class WenXinYiYanService:
    def __init__(self, api_endpoint):
        self.api_endpoint = api_endpoint

    def chatMessage(self, message):
        # 正确缩进的代码
        payload = json.dumps({
            "messages": [{
                "role": "user",
                "content": message
            }],
            "temperature": 0.95,
            "top_p": 0.8,
            "penalty_score": 1,
            "enable_system_memory": False,
            "disable_search": False,
            "enable_citation": False
        })
        headers = {'Content-Type': 'application/json'}
        response = requests.post(self.api_endpoint, headers=headers, data=payload)

        # 打印请求内容和响应内容
        print("Request Payload:", payload)
        print("Response Status Code:", response.status_code)

        try:
            # 检查返回的 JSON 格式
            response_json = response.json()  # 尝试将响应解析为 JSON
            print("Response JSON:", response_json)

            # 如果返回的 JSON 格式正确，尝试从中提取 'answer'
            return response_json.get('answer', '未能从文心一言API获取答案。')
        except ValueError as e:
            print(f"JSON解析错误：{e}")
            return "文心一言API返回的格式错误。"
        except Exception as e:
            print(f"发生错误：{e}")
            return f"请求文心一言API时发生错误：{e}"


class KimiService:
    def __init__(self, api_key, base_url, max_tokens=1024):  # 增加了max_tokens参数
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.max_tokens = max_tokens
    
    def chatMessage(self, message):
        try:
            completion = self.client.chat.completions.create(
                model="moonshot-v1-8k",
                messages=[{"role": "user", "content": message}],
                temperature=0.3,
                max_tokens=self.max_tokens  # 设置更大的max_tokens值
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"调用Kimi API时出错：{e}")
            return None

class AIAssistantDispatcher:
    def __init__(self):
        self.assistants = {
             'wenxinyiyan': WenXinYiYanService(api_endpoint="https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/completions_pro?access_token=24.d1ee8d6a70fcd162f7a26936f614650c.2592000.1736343757.282335-116580611"),
            'kimi': KimiService(api_key="sk-7H3ZdLK7dqilVMGQUUbtXQ1hrkf951tNYNR1TnEgNbG9gkGF", base_url="https://api.moonshot.cn/v1")
             }
        
        # 初始化BERT模型和tokenizer
        self.tokenizer = BertTokenizer.from_pretrained('C:/Users/12864/Desktop/model')
        self.model = BertModel.from_pretrained('C:/Users/12864/Desktop/model')
    
    def get_best_response(self, message, file_path=None):
        if file_path:
            # 如果有文件，就将文件内容作为消息的一部分
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    message += " " + content
            except Exception as e:
                return f"文件读取失败：{e}"

        # 使用并行请求多个AI助手生成不同部分的回答
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = {executor.submit(assistant.chatMessage, message): name for name, assistant in self.assistants.items()}
            responses = {}
            for future in concurrent.futures.as_completed(futures):
                name = futures[future]
                try:
                    response = future.result()
                    if response:
                        responses[name] = response
                except Exception as e:
                    responses[name] = f"调用{futures[future]}时发生错误：{e}"

        # 合并不同模型的回答
        merged_response = "\n".join(responses.values())

        # 目标函数：最大化文本的长度和相关性
        length_score = len(merged_response)
        relevance_score = self.calculate_relevance_score(merged_response, message)  # 计算相关性

        # 约束条件：确保回答长度不小于100个字符
        if length_score < 100:
            merged_response += " 这是为了增加回答的长度而添加的额外内容。"

        return merged_response if merged_response else "对不起，我无法回答您的问题。"
    
    def calculate_relevance_score(self, response, query):
        # 对问题和回答进行tokenization
        inputs_query = self.tokenizer(query, return_tensors='pt', padding=True, truncation=True, max_length=512)
        inputs_response = self.tokenizer(response, return_tensors='pt', padding=True, truncation=True, max_length=512)

        # 获取BERT的输出
        with torch.no_grad():
            query_output = self.model(**inputs_query)
            response_output = self.model(**inputs_response)

        # 提取[CLS] token的输出作为文本的表示
        query_embedding = query_output.last_hidden_state[:, 0, :]
        response_embedding = response_output.last_hidden_state[:, 0, :]

        # 计算余弦相似度
        similarity = torch.nn.functional.cosine_similarity(query_embedding, response_embedding)

        return similarity.item()

dispatcher = AIAssistantDispatcher()

# 定义默认路由，显示上传和提问的表单
@app.route('/')
def home():
    return render_template_string('''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>AI Assistant</title>
        <!-- 引入Bootstrap CSS -->
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
    </head>
    <body>
        <div class="container mt-5">
            <h1 class="text-center">AI Assistant</h1>
            <div class="row justify-content-center">
                <div class="col-md-6">
                    <form action="/upload" method="post" enctype="multipart/form-data" class="mb-3">
                        <div class="input-group">
                            <div class="custom-file">
                                <input type="file" class="custom-file-input" id="inputGroupFile01" name="file" required>
                                <label class="custom-file-label" for="inputGroupFile01">Choose file</label>
                            </div>
                            <div class="input-group-append">
                                <input type="submit" class="btn btn-primary" value="Upload and 解读 File">
                            </div>
                        </div>
                    </form>
                    <form action="/ask" method="post" class="mb-3">
                        <div class="input-group">
                            <input type="text" class="form-control" placeholder="Enter your question" name="question" required>
                            <div class="input-group-append">
                                <input type="submit" class="btn btn-outline-secondary" value="Ask a Question">
                            </div>
                        </div>
                    </form>
                </div>
            </div>
        </div>
        <!-- 引入Bootstrap JS -->
        <script src="https://code.jquery.com/jquery-3.5.1.slim.min.js"></script>
        <script src="https://cdn.jsdelivr.net/npm/@popperjs/core@2.9.3/dist/umd/popper.min.js"></script>
        <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/js/bootstrap.min.js"></script>
    </body>
    </html>
    ''')

# 文件上传路由
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        # 上传后直接解读文件
        best_answer = dispatcher.get_best_response("", file_path)
        
        # 将答案写入txt文件
        txt_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{filename}.txt')
        with open(txt_file_path, 'w', encoding='utf-8') as f:
            f.write(best_answer)
        
        # 将答案写入docx文件
        docx_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{filename}.docx')
        doc = Document()
        doc.add_paragraph(best_answer)
        doc.save(docx_file_path)
        
        # 返回下载链接
        return jsonify({
            'answer': best_answer,
            'txt_download_url': f'/downloads/{os.path.basename(txt_file_path)}',
            'docx_download_url': f'/downloads/{os.path.basename(docx_file_path)}'
        })

# 问题回答路由
@app.route('/ask', methods=['POST'])
def ask_question():
    user_question = request.form.get('question')
    if not user_question:
        return jsonify({'error': 'No question provided'}), 400
    best_answer = dispatcher.get_best_response(user_question)
    
    # 将答案写入txt文件
    txt_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{user_question.replace(" ", "_")}.txt')
    with open(txt_file_path, 'w', encoding='utf-8') as f:
        f.write(best_answer)
    
    # 将答案写入docx文件
    docx_file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f'answer_{user_question.replace(" ", "_")}.docx')
    doc = Document()
    doc.add_paragraph(best_answer)
    doc.save(docx_file_path)
    
    # 返回下载链接
    return jsonify({
        'answer': best_answer,
        'txt_download_url': f'/downloads/{os.path.basename(txt_file_path)}',
        'docx_download_url': f'/downloads/{os.path.basename(docx_file_path)}'
    })

# 下载文件路由
@app.route('/downloads/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(app.config['DOWNLOAD_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)
